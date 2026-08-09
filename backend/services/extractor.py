"""Text extraction services for PDF, TXT, and DOCX files."""

import re
from abc import ABC, abstractmethod
from typing import Any, Dict, List, Tuple
import pypdf
import docx

from models.document import DocumentPage


def clean_text(text: str) -> str:
    """Perform conservative text normalization without altering content."""
    if not text:
        return ""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove duplicate blank lines (>2 newlines reduced to 2)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim leading/trailing whitespace
    return text.strip()


class BaseExtractor(ABC):
    """Abstract base class for document text extractors."""

    @abstractmethod
    def extract(
        self, file_path: str
    ) -> Tuple[List[DocumentPage], Dict[str, Any]]:
        """Extract pages and metadata from the given file path."""
        pass


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents using pypdf."""

    def extract(
        self, file_path: str
    ) -> Tuple[List[DocumentPage], Dict[str, Any]]:
        try:
            reader = pypdf.PdfReader(file_path)
        except Exception as e:
            raise ValueError(f"Corrupted or invalid PDF file: {str(e)}")

        pages: List[DocumentPage] = []
        total_pages = len(reader.pages)

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            cleaned = clean_text(page_text)
            if cleaned:
                pages.append(DocumentPage(page_number=i + 1, text=cleaned))

        if not pages:
            raise ValueError("PDF contains no extractable text.")

        metadata: Dict[str, Any] = {
            "total_pages": total_pages,
            "extracted_pages": len(pages),
        }
        return pages, metadata


class TXTExtractor(BaseExtractor):
    """Extractor for plain text documents."""

    def extract(
        self, file_path: str
    ) -> Tuple[List[DocumentPage], Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
            except Exception as e:
                raise ValueError(f"Invalid text encoding: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")

        cleaned = clean_text(content)
        if not cleaned:
            raise ValueError("Text file is empty or contains only whitespace.")

        pages = [DocumentPage(page_number=None, text=cleaned)]
        metadata: Dict[str, Any] = {"encoding": "utf-8"}
        return pages, metadata


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX documents using python-docx."""

    def extract(
        self, file_path: str
    ) -> Tuple[List[DocumentPage], Dict[str, Any]]:
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise ValueError(f"Corrupted or invalid DOCX file: {str(e)}")

        paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraph_texts)
        cleaned = clean_text(full_text)

        if not cleaned:
            raise ValueError("DOCX contains no extractable text.")

        pages = [DocumentPage(page_number=None, text=cleaned)]
        metadata: Dict[str, Any] = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        return pages, metadata


def get_extractor(file_type: str) -> BaseExtractor:
    """Factory function to get the appropriate extractor for a file type."""
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return PDFExtractor()
    elif file_type == "txt":
        return TXTExtractor()
    elif file_type in ("docx", "doc"):
        return DOCXExtractor()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
